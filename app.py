import streamlit as st
import fitz  # PyMuPDF
from fpdf import FPDF
from docx import Document
import os
import tempfile
import base64
import requests

# ---- GROQ SETUP ----
GROQ_API_KEY = os.getenv("GROQ_API_KEY")  # Store securely or use st.secrets
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama3-70b-8192"  # Can be llama3-8b-8192 for lighter version

# ---- FUNCTIONS ----

def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text.strip()

def analyze_contract_with_groq(text):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": "You are a legal contract analyst."},
            {"role": "user", "content": f"Analyze the following contract and summarize its key terms, risks, and obligations:\n\n{text}"}
        ]
    }

    response = requests.post(GROQ_API_URL, headers=headers, json=payload)
    result = response.json()
    return result['choices'][0]['message']['content']

def generate_pdf_report(content, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.multi_cell(0, 10, line)
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    pdf.output(temp_path)
    return temp_path

def generate_docx_report(content, filename):
    doc = Document()
    doc.add_heading("Contract Analysis Report", 0)
    for para in content.split("\n"):
        doc.add_paragraph(para)
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(temp_path)
    return temp_path

# ---- STREAMLIT UI ----

st.title("📄 AI Contract Analyzer")
st.markdown("Upload a contract in PDF format to receive an AI-generated summary and risk analysis.")

uploaded_file = st.file_uploader("Upload Contract PDF", type=["pdf"])

output_format = st.selectbox("Choose Output Format", ["PDF", "Word (DOCX)"])

if uploaded_file and st.button("Analyze Contract"):
    with st.spinner("Extracting text and analyzing..."):
        text = extract_text_from_pdf(uploaded_file)
        if not text:
            st.error("No text found in PDF.")
        else:
            analysis = analyze_contract_with_groq(text)

            st.success("Analysis complete!")
            st.subheader("AI Summary:")
            st.text_area("Summary", analysis, height=300)

            filename = uploaded_file.name.replace(".pdf", f"_summary.{ 'pdf' if output_format == 'PDF' else 'docx'}")

            if output_format == "PDF":
                file_path = generate_pdf_report(analysis, filename)
            else:
                file_path = generate_docx_report(analysis, filename)

            with open(file_path, "rb") as f:
                st.download_button(
                    label=f"📥 Download {output_format} Report",
                    data=f,
                    file_name=filename,
                    mime="application/octet-stream"
                )
